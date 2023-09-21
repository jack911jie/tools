import os
import sys
from docx import Document
from docx.shared import Inches

class Picture:
    def __init__(self):
        pass

    def creat_docx(self,out_dir,title,pic_dir):
        # 创建一个新的文档
        doc = Document()

        # 添加标题
        if title=='':
            title='示例文档'
        doc.add_heading(title, level=1)

        # 插入图片
        for fn in os.listdir(pic_dir):
            if fn[-3:].lower()=='jpg':
                real_fn=os.path.join(pic_dir,fn)  

                doc.add_picture(real_fn, width=Inches(4.0))  # 插入图片并指定宽度
                # 添加文本
                txt=f'编号：{fn[:-4]}\n类型：\n价格：\n\n'
                doc.add_paragraph(txt)

        # 保存文档
        save_fn=os.path.join(out_dir,'首饰价格标注.docx')
        doc.save(save_fn)

        print('完成')


if __name__=='__main__':
    p=Picture()
    p.creat_docx(out_dir='E:\\temp\\ejj\\壹涧小铸首饰图\\2023020于二发饰',
                title='首饰价格',
                pic_dir='E:\\temp\\ejj\\壹涧小铸首饰图\\2023020于二发饰\\壹涧小铸首饰\\微文档照片')